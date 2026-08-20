from io import BytesIO
from unittest.mock import patch

import fitz
from docx import Document
from django.test import SimpleTestCase
from rest_framework.test import APIClient
from django.core.files.uploadedfile import SimpleUploadedFile

from .jd_parser import (
	UnsupportedJobDescriptionFormatError,
	parse_job_description,
)
from .resume_parser import UnsupportedResumeFormatError, parse_resume
from .skill_extractor import extract_keywords, extract_skills
from .similarity import (
	 screening_score,
	 semantic_similarity,
	 skill_coverage,
	 text_similarity,
)
from .ranking import rank_candidates


class ResumeParserTests(SimpleTestCase):
	def test_parses_pdf_resume(self):
		pdf = fitz.open()
		pdf.new_page().insert_text((72, 72), "Jane Doe\nPython Django")
		resume = BytesIO(pdf.tobytes())
		resume.name = "resume.pdf"

		text = parse_resume(resume)

		self.assertIn("Jane Doe", text)
		self.assertIn("Python Django", text)

	def test_parses_docx_resume(self):
		document = Document()
		document.add_paragraph("Jane Doe")
		document.add_paragraph("Python Django")
		resume = BytesIO()
		document.save(resume)
		resume.seek(0)
		resume.name = "resume.docx"

		self.assertEqual(parse_resume(resume), "Jane Doe\nPython Django")

	def test_rejects_legacy_doc_resume(self):
		resume = BytesIO()
		resume.name = "resume.doc"

		with self.assertRaises(UnsupportedResumeFormatError):
			parse_resume(resume)


class JobDescriptionParserTests(SimpleTestCase):
	def test_accepts_plain_text(self):
		text = parse_job_description("  Python developer with Django experience.  ")

		self.assertEqual(text, "Python developer with Django experience.")

	def test_parses_pdf_job_description(self):
		pdf = fitz.open()
		pdf.new_page().insert_text((72, 72), "Build APIs with Django")
		description = BytesIO(pdf.tobytes())
		description.name = "job-description.pdf"

		self.assertIn("Build APIs with Django", parse_job_description(description))

	def test_rejects_unsupported_format(self):
		description = BytesIO(b"role")
		description.name = "job-description.doc"

		with self.assertRaises(UnsupportedJobDescriptionFormatError):
			parse_job_description(description)


class SkillExtractorTests(SimpleTestCase):
	def test_extracts_normalized_skills_in_document_order(self):
		text = "We need Python, DRF, and machine learning. Python is required."

		self.assertEqual(
			extract_skills(text),
			["Python", "Django REST Framework", "Machine Learning"],
		)

	def test_does_not_match_skills_inside_words(self):
		self.assertEqual(extract_skills("A playful designer uses a stylesheet."), [])

	def test_extracts_custom_keywords(self):
		text = "The role supports FastAPI and PostgreSQL deployments."

		self.assertEqual(
			extract_keywords(text, ["FastAPI", "PostgreSQL"]),
			["FastAPI", "PostgreSQL"],
		)


class SimilarityTests(SimpleTestCase):
	def test_identical_text_has_full_similarity(self):
		text = "Python Django REST API development"

		self.assertAlmostEqual(text_similarity(text, text), 1.0)

	def test_unrelated_text_has_lower_similarity(self):
		self.assertLess(
			text_similarity("Python Django developer", "Graphic design illustrator"),
			0.5,
		)

	@patch("screening.similarity._get_embedding_model")
	def test_semantic_similarity_uses_embedding_model(self, get_model):
		class FakeModel:
			def encode(self, texts, normalize_embeddings):
				self.texts = texts
				self.normalize_embeddings = normalize_embeddings
				return [[1.0, 0.0], [1.0, 0.0]]

		get_model.return_value = FakeModel()

		self.assertEqual(semantic_similarity("Python developer", "API engineer"), 1.0)
		get_model.assert_called_once_with()

	def test_skill_coverage(self):
		resume = "Python and Django experience"
		job_description = "Need Python, Django, and Docker experience"

		self.assertAlmostEqual(skill_coverage(resume, job_description), 2 / 3)

	def test_screening_score_is_percentage(self):
		resume = "Python Django developer"
		job_description = "Python Django developer"

		self.assertEqual(screening_score(resume, job_description), 100.0)

	def test_rejects_zero_weights(self):
		with self.assertRaises(ValueError):
			screening_score("resume", "job", text_weight=0, skill_weight=0)


class RankingTests(SimpleTestCase):
	def test_ranks_candidates_highest_score_first(self):
		candidates = [
			{"name": "Partial", "resume_text": "Python experience"},
			{"name": "Best", "resume_text": "Python Django developer"},
		]

		results = rank_candidates(candidates, "Python Django developer")

		self.assertEqual([result["candidate"]["name"] for result in results], ["Best", "Partial"])
		self.assertEqual([result["rank"] for result in results], [1, 2])
		self.assertGreater(results[0]["score"], results[1]["score"])

	def test_preserves_order_for_tied_candidates(self):
		candidates = [
			{"name": "First", "resume_text": "Python"},
			{"name": "Second", "resume_text": "Python"},
		]

		results = rank_candidates(candidates, "Python")

		self.assertEqual([result["candidate"]["name"] for result in results], ["First", "Second"])

	def test_supports_object_candidates_with_custom_resume_getter(self):
		class Candidate:
			def __init__(self, name, resume):
				self.name = name
				self.resume = resume

		candidates = [Candidate("Best", "Django Python"), Candidate("Other", "Excel")]

		results = rank_candidates(
			candidates,
			"Django Python",
			resume_text_getter=lambda candidate: candidate.resume,
		)

		self.assertEqual(results[0]["candidate"].name, "Best")


class ScreeningAPITests(SimpleTestCase):
	def setUp(self):
		self.client = APIClient()

	@patch("screening.views.rank_candidates")
	def test_forwards_semantic_option(self, rank_candidates):
		pdf = fitz.open()
		pdf.new_page().insert_text((72, 72), "Python Django developer")
		resume = SimpleUploadedFile(
			"jane.pdf", pdf.tobytes(), content_type="application/pdf"
		)
		rank_candidates.return_value = [
			{
				"rank": 1,
				"score": 84.38,
				"candidate": {"filename": "jane.pdf"},
			}
		]

		response = self.client.post(
			"/api/screen/",
			{
				"job_description": "Python Django developer",
				"resumes": [resume],
				"use_semantic": "true",
			},
			format="multipart",
		)

		self.assertEqual(response.status_code, 200)
		rank_candidates.assert_called_once_with(
			[{
				"filename": "jane.pdf",
				"resume_text": "Python Django developer",
			}],
			"Python Django developer",
			use_semantic=True,
		)

	def test_screens_uploaded_resumes(self):
		pdf = fitz.open()
		pdf.new_page().insert_text((72, 72), "Python Django developer")
		resume = SimpleUploadedFile(
			"jane.pdf", pdf.tobytes(), content_type="application/pdf"
		)

		response = self.client.post(
			"/api/screen/",
			{"job_description": "Python Django developer", "resumes": [resume]},
			format="multipart",
		)

		self.assertEqual(response.status_code, 200)
		self.assertEqual(response.data["results"][0]["filename"], "jane.pdf")
		self.assertEqual(response.data["results"][0]["rank"], 1)

	def test_requires_job_description_and_resume(self):
		response = self.client.post("/api/screen/", {}, format="json")

		self.assertEqual(response.status_code, 400)
